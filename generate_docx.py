from docx import Document
from docx.shared import Pt
import os

doc = Document()
doc.add_heading('INF 406: E-COMMERCE II - Assignment Solutions', 0)

# Question 1
doc.add_heading('Question 1: E-commerce Website Security', level=1)
doc.add_paragraph('Prompt: Imagine you are the owner of an e-commerce website. What are some of the signs that your site has been hacked? Discuss the major types of attacks you could expect to experience and the resulting damage to your site. Prepare a brief summary presentation. [10 Marks]').style = 'Quote'

doc.add_heading('Part A: Signs Your E-commerce Site Has Been Hacked', level=2)
doc.add_paragraph('As an e-commerce owner, vigilance is key. Here are the primary indicators of a compromised site:')
signs = [
    ("Unexpected Content Changes: Unexplained modifications, such as defacement of the homepage, strange pop-ups, or hidden spam links injected into product pages."),
    ("Customer Complaints: An influx of users reporting that their credit card details were stolen shortly after purchasing from the site, or customers receiving phishing emails spoofing the company."),
    ("Performance Degradation: Unusually slow page load times, frequent server crashes, or unexplained bandwidth spikes, which may indicate a DDoS attack or a malware infection mining cryptocurrency."),
    ("Suspicious Admin Activity: The sudden appearance of new, unrecognized administrative accounts or unexpected password resets for existing staff."),
    ("Search Engine Warnings: Google or other search engines displaying 'This site may be hacked' or blacklisting the website entirely, leading to a massive drop in organic traffic."),
    ("Security Alerts: Warnings from web application firewalls (WAF), antivirus software, or file integrity monitoring systems indicating unauthorized file modifications.")
]
for sign in signs:
    doc.add_paragraph(sign, style='List Number')

doc.add_heading('Part B: Major Types of Attacks & Resulting Damage', level=2)
attacks = [
    ("Digital Skimming (Magecart Attacks)", "Mechanism: Hackers inject malicious JavaScript into the checkout pages to secretly capture customers' credit card information as they type it in.", "Damage: Devastating financial liability, regulatory fines (PCI-DSS violations), and catastrophic loss of customer trust."),
    ("SQL Injection (SQLi)", "Mechanism: Attackers insert malicious SQL statements into input fields (like search bars or login forms) to manipulate the backend database.", "Damage: Exposure of sensitive customer databases (passwords, personal info), data manipulation, or complete database deletion."),
    ("Distributed Denial of Service (DDoS)", "Mechanism: Overwhelming the e-commerce servers with a flood of fake traffic from a botnet, exhausting server resources.", "Damage: Extended website downtime, resulting in thousands of dollars in lost sales and frustrated customers abandoning their shopping carts."),
    ("Cross-Site Scripting (XSS)", "Mechanism: Injecting malicious scripts into web pages viewed by other users to steal session cookies.", "Damage: Attackers can hijack active customer sessions, potentially making unauthorized purchases or changing account details."),
    ("Ransomware", "Mechanism: Malware that encrypts the website's critical files or databases, demanding a ransom payment for the decryption key.", "Damage: Complete operational paralysis, severe financial loss, and potential permanent data loss if backups are compromised.")
]
for attack in attacks:
    doc.add_heading(attack[0], level=3)
    doc.add_paragraph(attack[1], style='List Bullet')
    doc.add_paragraph(attack[2], style='List Bullet')

# Question 2
doc.add_heading('Question 2: Certification Authorities (CAs)', level=1)
doc.add_paragraph('Prompt: Find three certification authorities and compare the features of each company’s digital certificates. Provide a brief description of each company as well, including number of clients. Prepare a brief presentation of your findings. [10 Marks]').style = 'Quote'

cas = [
    ("1. DigiCert", 
     "Description: DigiCert is the premier global provider of high-assurance digital certificates and enterprise-grade PKI (Public Key Infrastructure). They acquired Symantec's Website Security business.", 
     "Clients: Trusted by over 89% of Fortune 500 companies, securing millions of users globally.", 
     ["Offers premium Extended Validation (EV), Organization Validation (OV), and Domain Validation (DV) certificates.", "Provides advanced features like the DigiCert Smart Seal, Post-Quantum Cryptography (PQC) solutions, and high-tier warranties.", "Robust enterprise certificate lifecycle management via the DigiCert ONE platform."]),
    ("2. Sectigo (formerly Comodo CA)", 
     "Description: Sectigo is one of the largest and most well-known commercial CAs globally. It provides a wide range of web security products tailored for businesses of all sizes.", 
     "Clients: Serves over 700,000 businesses worldwide and issues tens of millions of active certificates.", 
     ["Comprehensive offerings including DV, OV, EV, Wildcard, and Multi-Domain (SAN) certificates.", "Standout features include daily website vulnerability scanning and automated malware removal tools bundled with premium certificates.", "Highly competitive pricing structure compared to DigiCert."]),
    ("3. Let's Encrypt", 
     "Description: Let's Encrypt is a free, automated, and open Certificate Authority run by the Internet Security Research Group (ISRG), a non-profit organization.", 
     "Clients: Secures over 300 million websites globally, making it the largest CA in the world by sheer volume.", 
     ["100% Free: Does not charge for any of its certificates.", "Automation-First: Certificates are valid for only 90 days, strongly encouraging users to use ACME clients to automate the renewal process.", "Validation Types: Only offers Domain Validation (DV) and Wildcard certificates. It does not provide OV or EV certificates."])
]

for ca in cas:
    doc.add_heading(ca[0], level=2)
    doc.add_paragraph(ca[1], style='List Bullet')
    doc.add_paragraph(ca[2], style='List Bullet')
    p = doc.add_paragraph('Certificate Features:')
    for feat in ca[3]:
        doc.add_paragraph(feat, style='List Continue')
        
doc.add_heading('Brief Comparison Summary', level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'DigiCert'
hdr_cells[2].text = 'Sectigo'
hdr_cells[3].text = 'Let\'s Encrypt'

data = [
    ('Target Audience', 'Large Enterprises / Fortune 500', 'SMBs to Enterprises', 'Personal Sites, Developers, General Web'),
    ('Cost', 'Premium / High', 'Affordable / Mid-range', '100% Free'),
    ('Validation Offered', 'DV, OV, EV', 'DV, OV, EV', 'DV only'),
    ('Lifespan', 'Up to 1 year', 'Up to 1 year', '90 Days (Automated)'),
    ('Warranty', 'Up to $2,000,000', 'Up to $1,750,000', 'None'),
    ('Key Strength', 'Highest trust, enterprise management', 'Bundled security tools, flexible pricing', 'Free, universally accessible, automated')
]

for row in data:
    row_cells = table.add_row().cells
    for i in range(4):
        row_cells[i].text = row[i]

# References
doc.add_heading('References', level=1)
refs = [
    "OWASP Foundation. 'SQL Injection.' https://owasp.org/www-community/attacks/SQL_Injection",
    "Cloudflare. 'What is a Magecart attack?' https://www.cloudflare.com/learning/security/what-is-magecart/",
    "DigiCert. 'Digital Certificates: What Are They and How Do They Work?' https://www.digicert.com/tls-ssl/digital-certificate",
    "Sectigo. 'About Sectigo: Digital Identity Management.' https://sectigo.com/about",
    "Let's Encrypt. 'About Let's Encrypt.' Internet Security Research Group (ISRG). https://letsencrypt.org/about/"
]
for ref in refs:
    doc.add_paragraph(ref, style='List Bullet')

output_path = r"C:\Users\PC-TECH\Desktop\L400 Sem 2\E-Commerce_II_Assignment_Solutions_with_References.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
